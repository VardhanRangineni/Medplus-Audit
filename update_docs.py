#!/usr/bin/env python3
"""
Update documentation content in DOCX files while preserving structure
"""
from docx import Document
import re

def update_auditor_modal_docx(doc_path):
    """Update auditor modal content in DOCX file"""
    doc = Document(doc_path)
    
    replacements = {
        # Auditor identification
        'Abhinay Naidu': 'Srikanth Rao',
        'ID: A046': 'ID: A039',
        
        # Performance cards
        'Number: 37': 'Number: 26',
        '37 (total audits completed': '26 (total audits completed',
        'Number: 45,389': 'Number: 31,414',
        '45,389 (Physical Inventory': '31,414 (Physical Inventory',
        'Number: 1.81 L': 'Number: 1.29 L',
        '1.81 L (181,000 Stock': '1.29 L (Stock',
        '181,000 Stock Keeping Units': 'Stock Keeping Units',
        
        # Deviation summary - Appeared
        'SKUs: 13,081': 'SKUs: 10,017',
        'Qty: 2.05 L': 'Qty: 1.52 L',
        'Value: ₹1.33 Cr': 'Value: ₹1.09 Cr',
        
        # Deviation summary - Matched
        'SKUs: 12,221': 'SKUs: 9,345',
        'Qty: 1.92 L': 'Qty: 1.42 L',
        'Value: ₹1.25 Cr': 'Value: ₹1.03 Cr',
        
        # Deviation summary - Revised
        'SKUs: 860': 'SKUs: 672',
        'Qty: 13,615': 'Qty: 9,770',
        'Value: ₹8.51 L': 'Value: ₹6.53 L',
    }
    
    # Update paragraphs
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in para.text:
                # Replace in runs to preserve formatting
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
    
    # Update tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in para.text:
                            for run in para.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
    
    doc.save(doc_path)
    print(f"Updated: {doc_path}")

def update_supervisor_modal_docx(doc_path):
    """Update supervisor modal content in DOCX file"""
    doc = Document(doc_path)
    
    replacements = {
        # Supervisor identification
        'SUP001': 'S001',
        
        # Add example name
        'Display name of the supervisor': 'Display name of the supervisor (e.g., "Aditya Reddy")',
        'Unique identifier (e.g., SUP001)': 'Unique identifier (e.g., "S001")',
        
        # Update card count description
        'Four primary KPI cards': 'Six primary KPI cards',
        
        # Add specific values for the cards
        'Count of unique audits supervised': 'Count: 15 (unique audits supervised',
        'Number of distinct days the supervisor': 'Count: 60 (distinct days',
        'Aggregate count of Product IDs': 'Count: 70,251 (Product IDs',
        'Total Stock Keeping Units supervised': 'Count: 3.40 L (Stock Keeping Units supervised)',
        
        # Deviation values
        'Quantity: Total count of items with appeared deviations': 'SKUs: 32,149\n- Qty: 5.73 L\n- Value: ₹2.13 Cr\n- Total count of items with appeared deviations',
        'Quantity: Items where initial deviation was confirmed': 'SKUs: 29,953\n- Qty: 5.34 L\n- Value: ₹1.98 Cr\n- Items where initial deviation was confirmed',
        'Quantity: Items that required correction by supervisor': 'SKUs: 2,196\n- Qty: 38,456\n- Value: ₹14.90 L\n- Items that required correction by supervisor',
    }
    
    # Update paragraphs
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in para.text:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
    
    # Update tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in para.text:
                            for run in para.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
    
    doc.save(doc_path)
    print(f"Updated: {doc_path}")

def update_combined_docx(doc_path):
    """Update both auditor and supervisor sections in combined DOCX"""
    doc = Document(doc_path)
    
    all_replacements = {
        # Auditor section
        'Abhinay Naidu': 'Srikanth Rao',
        'ID: A046': 'ID: A039',
        'Number: 37': 'Number: 26',
        '37 (total audits completed': '26 (total audits completed',
        'Number: 45,389': 'Number: 31,414',
        '45,389 (Physical Inventory': '31,414 (Physical Inventory',
        'Number: 1.81 L': 'Number: 1.29 L',
        '1.81 L (181,000 Stock': '1.29 L (Stock',
        'SKUs: 13,081': 'SKUs: 10,017',
        'Qty: 2.05 L': 'Qty: 1.52 L',
        'Value: ₹1.33 Cr': 'Value: ₹1.09 Cr',
        'SKUs: 12,221': 'SKUs: 9,345',
        'Qty: 1.92 L': 'Qty: 1.42 L',
        'Value: ₹1.25 Cr': 'Value: ₹1.03 Cr',
        'SKUs: 860': 'SKUs: 672',
        'Qty: 13,615': 'Qty: 9,770',
        'Value: ₹8.51 L': 'Value: ₹6.53 L',
        
        # Supervisor section
        'SUP001': 'S001',
        'Four primary KPI cards': 'Six primary KPI cards',
    }
    
    # Update all content
    for para in doc.paragraphs:
        for old_text, new_text in all_replacements.items():
            if old_text in para.text:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in all_replacements.items():
                        if old_text in para.text:
                            for run in para.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
    
    doc.save(doc_path)
    print(f"Updated: {doc_path}")

if __name__ == "__main__":
    import os
    docs_dir = "/home/runner/work/Medplus-Audit/Medplus-Audit/docs"
    
    # Update individual DOCX files if they exist
    auditor_docx = os.path.join(docs_dir, "auditor-detail-modal.docx")
    if os.path.exists(auditor_docx):
        update_auditor_modal_docx(auditor_docx)
    
    supervisor_docx = os.path.join(docs_dir, "supervisor-modal.docx")
    if os.path.exists(supervisor_docx):
        update_supervisor_modal_docx(supervisor_docx)
    
    # Update combined documentation
    combined_docx = os.path.join(docs_dir, "COMBINED_DOCUMENTATION.docx")
    if os.path.exists(combined_docx):
        update_combined_docx(combined_docx)
    
    print("All DOCX files updated successfully!")
